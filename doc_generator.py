"""Generacion de reportes Excel/PDF en memoria (bytes) para Streamlit.

Versiones simplificadas de los reportes de scripts/reportes.py, adaptadas
para retornar bytes en lugar de escribir archivos a disco.
"""

from io import BytesIO
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
)

from db import query, fmt_money


# ---------------------------------------------------------------------------
# Helpers Excel
# ---------------------------------------------------------------------------

VERDE = "2E5C3F"
CREMA = "FDF8F0"
SECUNDARIO = "EFF5ED"

HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
HEADER_FILL = PatternFill("solid", fgColor=VERDE)
TITLE_FONT = Font(bold=True, size=14, color=VERDE)
BORDER = Border(
    left=Side(style="thin", color="CCCCCC"),
    right=Side(style="thin", color="CCCCCC"),
    top=Side(style="thin", color="CCCCCC"),
    bottom=Side(style="thin", color="CCCCCC"),
)


def _write_headers(ws, row, headers):
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=row, column=col, value=h)
        c.font = HEADER_FONT
        c.fill = HEADER_FILL
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = BORDER


def _write_rows(ws, start_row, data, money_cols=None):
    money_cols = money_cols or []
    for i, row_data in enumerate(data):
        for col, val in enumerate(row_data, 1):
            c = ws.cell(row=start_row + i, column=col, value=val)
            c.border = BORDER
            if col in money_cols and isinstance(val, (int, float)):
                c.number_format = '"$"#,##0.00'


def _auto_width(ws):
    for col in ws.columns:
        max_len = 0
        letter = None
        for cell in col:
            if hasattr(cell, "column_letter") and letter is None:
                letter = cell.column_letter
            if cell.value is not None:
                max_len = max(max_len, len(str(cell.value)))
        if letter:
            ws.column_dimensions[letter].width = min(max_len + 2, 40)


# ---------------------------------------------------------------------------
# Reportes Excel
# ---------------------------------------------------------------------------

def generar_reporte_mensual_bytes(mes):
    """Reporte mensual consolidado (cobranza, gastos, P&L) como bytes XLSX."""
    wb = Workbook()

    # Resumen
    ws = wb.active
    ws.title = "Resumen"
    ws["A1"] = f"Reporte Mensual — {mes}"
    ws["A1"].font = TITLE_FONT
    ws.merge_cells("A1:D1")

    facturado = query(
        "SELECT COALESCE(SUM(monto), 0) AS t FROM cargos WHERE mes_aplicable = ?",
        (mes,), fetchone=True,
    )["t"] or 0
    cobrado = query(
        """SELECT COALESCE(SUM(p.monto_pagado), 0) AS t FROM pagos p
           JOIN cargos c ON c.id = p.cargo_id WHERE c.mes_aplicable = ?""",
        (mes,), fetchone=True,
    )["t"] or 0
    gastos = query(
        "SELECT COALESCE(SUM(monto), 0) AS t FROM gastos WHERE substr(fecha, 1, 7) = ?",
        (mes,), fetchone=True,
    )["t"] or 0
    nomina = query(
        "SELECT COALESCE(SUM(monto), 0) AS t FROM pagos_nomina WHERE periodo = ?",
        (mes,), fetchone=True,
    )["t"] or 0

    _write_headers(ws, 3, ["Concepto", "Monto"])
    _write_rows(ws, 4, [
        ("Ingresos facturados", facturado),
        ("Ingresos cobrados", cobrado),
        ("Pendiente por cobrar", facturado - cobrado),
        ("Gastos operativos", gastos),
        ("Nomina", nomina),
        ("Utilidad neta", cobrado - gastos - nomina),
    ], money_cols=[2])
    _auto_width(ws)

    # Cobranza
    ws2 = wb.create_sheet("Cobranza")
    ws2["A1"] = f"Cobranza — {mes}"
    ws2["A1"].font = TITLE_FONT
    ws2.merge_cells("A1:F1")
    _write_headers(ws2, 3, ["Familia", "Alumno", "Concepto", "Monto", "Pagado", "Estatus"])
    cargos = query(
        """SELECT f.nombre_familia AS familia, a.nombre AS alumno,
                  co.nombre AS concepto, c.monto AS monto,
                  COALESCE((SELECT SUM(monto_pagado) FROM pagos WHERE cargo_id = c.id), 0) AS pagado,
                  c.estatus AS estatus
           FROM cargos c
           JOIN alumnos a ON a.id = c.alumno_id
           JOIN familias f ON f.id = a.familia_id
           JOIN conceptos_cobro co ON co.id = c.concepto_id
           WHERE c.mes_aplicable = ?
           ORDER BY f.nombre_familia""",
        (mes,),
    )
    _write_rows(ws2, 4, [
        (r["familia"], r["alumno"], r["concepto"], r["monto"], r["pagado"], r["estatus"])
        for r in cargos
    ], money_cols=[4, 5])
    _auto_width(ws2)

    # Gastos
    ws3 = wb.create_sheet("Gastos")
    ws3["A1"] = f"Gastos — {mes}"
    ws3["A1"].font = TITLE_FONT
    ws3.merge_cells("A1:E1")
    _write_headers(ws3, 3, ["Fecha", "Categoria", "Descripcion", "Proveedor", "Monto"])
    gastos_rows = query(
        """SELECT g.fecha, cg.nombre AS categoria, g.descripcion,
                  COALESCE(p.nombre, '') AS proveedor, g.monto
           FROM gastos g
           LEFT JOIN categorias_gasto cg ON cg.id = g.categoria_id
           LEFT JOIN proveedores p ON p.id = g.proveedor_id
           WHERE substr(g.fecha, 1, 7) = ?
           ORDER BY g.fecha""",
        (mes,),
    )
    _write_rows(ws3, 4, [
        (r["fecha"], r["categoria"], r["descripcion"], r["proveedor"], r["monto"])
        for r in gastos_rows
    ], money_cols=[5])
    _auto_width(ws3)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def generar_reporte_contador_bytes(mes):
    """Reporte para contador (formato fiscal simplificado)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Ingresos"
    ws["A1"] = f"Ingresos — {mes}"
    ws["A1"].font = TITLE_FONT
    ws.merge_cells("A1:E1")
    _write_headers(ws, 3, ["Fecha", "Familia", "Concepto", "Monto", "Metodo"])
    ingresos = query(
        """SELECT p.fecha_pago AS fecha, f.nombre_familia AS familia,
                  co.nombre AS concepto,
                  p.monto_pagado AS monto, p.metodo_pago AS metodo
           FROM pagos p
           JOIN cargos c ON c.id = p.cargo_id
           JOIN alumnos a ON a.id = c.alumno_id
           JOIN familias f ON f.id = a.familia_id
           JOIN conceptos_cobro co ON co.id = c.concepto_id
           WHERE substr(p.fecha_pago, 1, 7) = ?
           ORDER BY p.fecha_pago""",
        (mes,),
    )
    _write_rows(ws, 4, [
        (r["fecha"], r["familia"], r["concepto"], r["monto"], r["metodo"])
        for r in ingresos
    ], money_cols=[4])
    _auto_width(ws)

    ws2 = wb.create_sheet("Egresos")
    ws2["A1"] = f"Egresos — {mes}"
    ws2["A1"].font = TITLE_FONT
    ws2.merge_cells("A1:F1")
    _write_headers(ws2, 3, ["Fecha", "Categoria", "Descripcion", "Proveedor", "Comprobante", "Monto"])
    egresos = query(
        """SELECT g.fecha, cg.nombre AS categoria, g.descripcion,
                  COALESCE(p.nombre, '') AS proveedor,
                  COALESCE(g.comprobante_ref, '') AS comprobante, g.monto
           FROM gastos g
           LEFT JOIN categorias_gasto cg ON cg.id = g.categoria_id
           LEFT JOIN proveedores p ON p.id = g.proveedor_id
           WHERE substr(g.fecha, 1, 7) = ?
           ORDER BY g.fecha""",
        (mes,),
    )
    _write_rows(ws2, 4, [
        (r["fecha"], r["categoria"], r["descripcion"], r["proveedor"], r["comprobante"], r["monto"])
        for r in egresos
    ], money_cols=[6])
    _auto_width(ws2)

    ws3 = wb.create_sheet("Nomina")
    ws3["A1"] = f"Nomina — {mes}"
    ws3["A1"].font = TITLE_FONT
    ws3.merge_cells("A1:D1")
    _write_headers(ws3, 3, ["Empleado", "Puesto", "Fecha pago", "Monto"])
    nomina = query(
        """SELECT e.nombre, e.puesto, pn.fecha_pago AS fecha, pn.monto
           FROM pagos_nomina pn JOIN empleados e ON e.id = pn.empleado_id
           WHERE pn.periodo = ? ORDER BY e.nombre""",
        (mes,),
    )
    _write_rows(ws3, 4, [
        (r["nombre"], r["puesto"], r["fecha"], r["monto"])
        for r in nomina
    ], money_cols=[4])
    _auto_width(ws3)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Recibo PDF
# ---------------------------------------------------------------------------

def generar_recibo_bytes(pago_id):
    """Recibo PDF de un pago como bytes."""
    pago = query(
        """SELECT p.id, p.fecha_pago AS fecha, p.monto_pagado,
                  p.metodo_pago AS metodo, p.referencia,
                  c.monto AS monto_cargo, c.mes_aplicable,
                  a.nombre AS alumno, f.nombre_familia AS familia,
                  f.contacto_principal, co.nombre AS concepto
           FROM pagos p
           JOIN cargos c ON c.id = p.cargo_id
           JOIN alumnos a ON a.id = c.alumno_id
           JOIN familias f ON f.id = a.familia_id
           JOIN conceptos_cobro co ON co.id = c.concepto_id
           WHERE p.id = ?""",
        (pago_id,), fetchone=True,
    )
    if not pago:
        return None

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"],
                        textColor=colors.HexColor(f"#{VERDE}"), alignment=1)
    normal = styles["Normal"]

    story = []
    story.append(Paragraph("Cavalletto — Preescolar de Naturaleza", h1))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(f"<b>Recibo de pago #{pago['id']}</b>", normal))
    story.append(Paragraph(f"Fecha: {pago['fecha']}", normal))
    story.append(Spacer(1, 0.5*cm))

    data = [
        ["Familia", pago["familia"]],
        ["Contacto", pago["contacto_principal"]],
        ["Alumno", pago["alumno"]],
        ["Concepto", pago["concepto"]],
        ["Mes aplicable", pago["mes_aplicable"]],
        ["Monto del cargo", fmt_money(pago["monto_cargo"])],
        ["Monto pagado", fmt_money(pago["monto_pagado"])],
        ["Metodo", pago["metodo"] or ""],
        ["Referencia", pago["referencia"] or ""],
    ]
    tbl = Table(data, colWidths=[5*cm, 10*cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{SECUNDARIO}")),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor(f"#{VERDE}")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(
        f"<i>Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}</i>", normal))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Proceso Administrativo — Word (.docx)
# ---------------------------------------------------------------------------

def generar_proceso_docx_bytes(proceso):
    """Genera documento Word profesional de un proceso administrativo.

    Args:
        proceso: dict con nombre, objetivo, area, frecuencia, pasos, kpis, etc.
    Returns:
        bytes del .docx
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Estilos
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    verde_rgb = RGBColor(0x2E, 0x5C, 0x3F)

    # Titulo
    titulo = doc.add_heading(proceso.get("nombre", "Proceso"), level=0)
    titulo.runs[0].font.color.rgb = verde_rgb

    # Subtitulo
    sub = doc.add_paragraph()
    sub.add_run("Cavalletto - Preescolar de Naturaleza").bold = True
    sub.add_run(f"\nDocumento de Proceso Administrativo")

    doc.add_paragraph()

    # Ficha tecnica
    doc.add_heading("Ficha Tecnica", level=1)
    ficha_data = [
        ("Area", proceso.get("area", "")),
        ("Frecuencia", proceso.get("frecuencia", "")),
        ("Responsable principal", proceso.get("responsable_principal", "Por asignar")),
        ("Estatus", proceso.get("estatus", "borrador")),
        ("Version", str(proceso.get("version", 1))),
        ("Trigger de inicio", proceso.get("trigger_inicio", "N/A")),
    ]
    table = doc.add_table(rows=len(ficha_data), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, val) in enumerate(ficha_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True if table.rows[i].cells[0].paragraphs[0].runs else False

    doc.add_paragraph()

    # Objetivo
    doc.add_heading("Objetivo", level=1)
    doc.add_paragraph(proceso.get("objetivo", ""))

    # Pasos
    pasos = proceso.get("pasos", [])
    if pasos:
        doc.add_heading("Pasos del Proceso", level=1)
        for i, paso in enumerate(pasos, 1):
            accion = paso.get("accion", paso.get("nombre", "Sin nombre"))
            h = doc.add_heading(f"Paso {i}: {accion}", level=2)

            if paso.get("responsable"):
                p = doc.add_paragraph()
                p.add_run("Responsable: ").bold = True
                p.add_run(paso["responsable"])
            if paso.get("tiempo_estimado"):
                p = doc.add_paragraph()
                p.add_run("Tiempo estimado: ").bold = True
                p.add_run(paso["tiempo_estimado"])
            if paso.get("herramienta"):
                p = doc.add_paragraph()
                p.add_run("Herramienta/Sistema: ").bold = True
                p.add_run(paso["herramienta"])
            if paso.get("entregable"):
                p = doc.add_paragraph()
                p.add_run("Entregable: ").bold = True
                p.add_run(paso["entregable"])
            if paso.get("criterio_exito"):
                p = doc.add_paragraph()
                p.add_run("Criterio de exito: ").bold = True
                p.add_run(paso["criterio_exito"])
            if paso.get("notas"):
                p = doc.add_paragraph()
                p.add_run("Nota: ").italic = True
                p.add_run(paso["notas"])

    # KPIs
    kpis = proceso.get("kpis", [])
    if kpis:
        doc.add_heading("KPIs y Metricas", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Metrica"
        hdr[1].text = "Meta"
        hdr[2].text = "Medicion"
        for kpi in kpis:
            row = table.add_row().cells
            if isinstance(kpi, dict):
                row[0].text = kpi.get("nombre", "")
                row[1].text = kpi.get("meta", "")
                row[2].text = kpi.get("medicion", "")
            else:
                row[0].text = str(kpi)

    # Excepciones
    excepciones = proceso.get("excepciones", [])
    if excepciones:
        doc.add_heading("Excepciones y Reglas Especiales", level=1)
        for exc in excepciones:
            if isinstance(exc, dict):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"Si ").bold = True
                p.add_run(exc.get("condicion", ""))
                p.add_run(f" entonces ").bold = True
                p.add_run(exc.get("accion", ""))
            else:
                doc.add_paragraph(str(exc), style="List Bullet")

    # Automatizaciones
    automatizaciones = proceso.get("automatizaciones", [])
    if automatizaciones:
        doc.add_heading("Automatizaciones Vinculadas", level=1)
        for auto in automatizaciones:
            if isinstance(auto, dict):
                doc.add_paragraph(
                    f"{auto.get('script', '')} - {auto.get('comando', '')}",
                    style="List Bullet"
                )
            else:
                doc.add_paragraph(str(auto), style="List Bullet")

    # Notas
    if proceso.get("notas"):
        doc.add_heading("Notas Adicionales", level=1)
        doc.add_paragraph(proceso["notas"])

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(
        f"Generado por Agente Administrativo Cavalletto - {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ).italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Proceso Administrativo — Excel (.xlsx)
# ---------------------------------------------------------------------------

def generar_proceso_xlsx_bytes(proceso):
    """Genera Excel con el proceso: ficha + pasos + KPIs en hojas separadas.

    Args:
        proceso: dict con nombre, objetivo, area, frecuencia, pasos, kpis, etc.
    Returns:
        bytes del .xlsx
    """
    wb = Workbook()

    # Hoja 1: Ficha + Pasos
    ws = wb.active
    ws.title = "Proceso"
    ws["A1"] = proceso.get("nombre", "Proceso")
    ws["A1"].font = TITLE_FONT
    ws.merge_cells("A1:F1")

    # Ficha
    ws["A3"] = "Area"
    ws["B3"] = proceso.get("area", "")
    ws["A4"] = "Frecuencia"
    ws["B4"] = proceso.get("frecuencia", "")
    ws["A5"] = "Responsable"
    ws["B5"] = proceso.get("responsable_principal", "")
    ws["A6"] = "Trigger"
    ws["B6"] = proceso.get("trigger_inicio", "")
    ws["A7"] = "Estatus"
    ws["B7"] = proceso.get("estatus", "borrador")
    ws["A8"] = "Objetivo"
    ws["B8"] = proceso.get("objetivo", "")
    for r in range(3, 9):
        ws.cell(row=r, column=1).font = Font(bold=True, color=VERDE)

    # Pasos
    pasos = proceso.get("pasos", [])
    if pasos:
        row = 10
        _write_headers(ws, row, ["#", "Accion", "Responsable", "Tiempo", "Herramienta", "Entregable", "Criterio de exito"])
        for i, paso in enumerate(pasos, 1):
            row += 1
            data = [
                i,
                paso.get("accion", paso.get("nombre", "")),
                paso.get("responsable", ""),
                paso.get("tiempo_estimado", ""),
                paso.get("herramienta", ""),
                paso.get("entregable", ""),
                paso.get("criterio_exito", ""),
            ]
            for col, val in enumerate(data, 1):
                c = ws.cell(row=row, column=col, value=val)
                c.border = BORDER
    _auto_width(ws)

    # Hoja 2: KPIs
    kpis = proceso.get("kpis", [])
    if kpis:
        ws2 = wb.create_sheet("KPIs")
        ws2["A1"] = "KPIs del Proceso"
        ws2["A1"].font = TITLE_FONT
        _write_headers(ws2, 3, ["Metrica", "Meta", "Medicion"])
        for i, kpi in enumerate(kpis):
            row_data = []
            if isinstance(kpi, dict):
                row_data = [kpi.get("nombre", ""), kpi.get("meta", ""), kpi.get("medicion", "")]
            else:
                row_data = [str(kpi), "", ""]
            for col, val in enumerate(row_data, 1):
                c = ws2.cell(row=4 + i, column=col, value=val)
                c.border = BORDER
        _auto_width(ws2)

    # Hoja 3: Excepciones
    excepciones = proceso.get("excepciones", [])
    if excepciones:
        ws3 = wb.create_sheet("Excepciones")
        ws3["A1"] = "Excepciones y Reglas"
        ws3["A1"].font = TITLE_FONT
        _write_headers(ws3, 3, ["Condicion", "Accion"])
        for i, exc in enumerate(excepciones):
            if isinstance(exc, dict):
                ws3.cell(row=4 + i, column=1, value=exc.get("condicion", "")).border = BORDER
                ws3.cell(row=4 + i, column=2, value=exc.get("accion", "")).border = BORDER
            else:
                ws3.cell(row=4 + i, column=1, value=str(exc)).border = BORDER
        _auto_width(ws3)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Proceso Administrativo — PDF
# ---------------------------------------------------------------------------

def generar_proceso_pdf_bytes(proceso):
    """Genera PDF profesional de un proceso administrativo."""
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"],
                        textColor=colors.HexColor(f"#{VERDE}"), alignment=1)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"],
                        textColor=colors.HexColor(f"#{VERDE}"))
    normal = styles["Normal"]

    story = []
    story.append(Paragraph(proceso.get("nombre", "Proceso"), h1))
    story.append(Paragraph("Cavalletto - Preescolar de Naturaleza", normal))
    story.append(Spacer(1, 0.5*cm))

    # Ficha
    ficha = [
        ["Area", proceso.get("area", "")],
        ["Frecuencia", proceso.get("frecuencia", "")],
        ["Responsable", proceso.get("responsable_principal", "Por asignar")],
        ["Trigger", proceso.get("trigger_inicio", "N/A")],
        ["Estatus", proceso.get("estatus", "borrador")],
    ]
    tbl = Table(ficha, colWidths=[4*cm, 12*cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{SECUNDARIO}")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.5*cm))

    # Objetivo
    story.append(Paragraph("Objetivo", h2))
    story.append(Paragraph(proceso.get("objetivo", ""), normal))
    story.append(Spacer(1, 0.3*cm))

    # Pasos
    pasos = proceso.get("pasos", [])
    if pasos:
        story.append(Paragraph("Pasos del Proceso", h2))
        pasos_data = [["#", "Accion", "Responsable", "Tiempo", "Entregable"]]
        for i, p in enumerate(pasos, 1):
            pasos_data.append([
                str(i),
                p.get("accion", p.get("nombre", "")),
                p.get("responsable", ""),
                p.get("tiempo_estimado", ""),
                p.get("entregable", ""),
            ])
        tbl2 = Table(pasos_data, colWidths=[1*cm, 5*cm, 3.5*cm, 2.5*cm, 4*cm])
        tbl2.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{VERDE}")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{CREMA}")]),
        ]))
        story.append(tbl2)
        story.append(Spacer(1, 0.3*cm))

    # KPIs
    kpis = proceso.get("kpis", [])
    if kpis:
        story.append(Paragraph("KPIs", h2))
        kpi_data = [["Metrica", "Meta", "Medicion"]]
        for k in kpis:
            if isinstance(k, dict):
                kpi_data.append([k.get("nombre", ""), k.get("meta", ""), k.get("medicion", "")])
            else:
                kpi_data.append([str(k), "", ""])
        tbl3 = Table(kpi_data, colWidths=[5*cm, 5.5*cm, 5.5*cm])
        tbl3.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{VERDE}")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(tbl3)

    # Footer
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(
        f"<i>Generado por Agente Administrativo Cavalletto - {datetime.now().strftime('%d/%m/%Y %H:%M')}</i>",
        normal))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def generate(reporte, **kw):
    """Genera bytes segun tipo de reporte. Retorna (bytes, filename, mime) o None."""
    if reporte == "mensual":
        mes = kw.get("mes")
        data = generar_reporte_mensual_bytes(mes)
        return data, f"reporte_mensual_{mes}.xlsx", \
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if reporte == "contador":
        mes = kw.get("mes")
        data = generar_reporte_contador_bytes(mes)
        return data, f"reporte_contador_{mes}.xlsx", \
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if reporte == "recibo":
        pago_id = kw.get("pago_id")
        data = generar_recibo_bytes(pago_id)
        if data is None:
            return None
        return data, f"recibo_{pago_id}.pdf", "application/pdf"
    if reporte == "proceso_docx":
        proceso = kw.get("proceso")
        if not proceso:
            return None
        slug = proceso.get("nombre", "proceso").lower().replace(" ", "_")
        data = generar_proceso_docx_bytes(proceso)
        return data, f"{slug}.docx", \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if reporte == "proceso_xlsx":
        proceso = kw.get("proceso")
        if not proceso:
            return None
        slug = proceso.get("nombre", "proceso").lower().replace(" ", "_")
        data = generar_proceso_xlsx_bytes(proceso)
        return data, f"{slug}.xlsx", \
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if reporte == "proceso_pdf":
        proceso = kw.get("proceso")
        if not proceso:
            return None
        slug = proceso.get("nombre", "proceso").lower().replace(" ", "_")
        data = generar_proceso_pdf_bytes(proceso)
        return data, f"{slug}.pdf", "application/pdf"
    return None
